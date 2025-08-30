"""
Document Writer Tool - A utility for writing content to Word or Markdown files.
"""

import os
import pandas as pd
from typing import Optional, List, Dict, Any, Union
from docx import Document
from docx.shared import Inches


class DocumentHandler:
    """
    A class to write and read content from various document formats including Word and Markdown.
    """

    @staticmethod
    def read_word(filepath: str) -> str:
        """
        Read content from a Word file.
        
        Args:
            filepath (str): The path to the Word file to read from
            
        Returns:
            str: The content of the Word file as a string
            
        Raises:
            FileNotFoundError: If the file does not exist
            Exception: If there is an error reading the file
        """
        try:
            if not os.path.exists(filepath):
                raise FileNotFoundError(f"File {filepath} does not exist")
            
            doc = Document(filepath)
            content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                content.append(paragraph.text)
                
            # Join paragraphs with newlines
            return '\n'.join(content)
        except Exception as e:
            print(f"Error reading Word file: {e}")
            raise
    
    @staticmethod
    def read_markdown(filepath: str) -> str:
        """
        Read content from a Markdown file.
        
        Args:
            filepath (str): The path to the Markdown file to read from
            
        Returns:
            str: The content of the Markdown file as a string
            
        Raises:
            FileNotFoundError: If the file does not exist
            Exception: If there is an error reading the file
        """
        try:
            if not os.path.exists(filepath):
                raise FileNotFoundError(f"File {filepath} does not exist")
            
            with open(filepath, 'r', encoding='utf-8') as f:
                content = f.read()
            return content
        except Exception as e:
            print(f"Error reading Markdown file: {e}")
            raise

    @staticmethod
    def read_csv(filepath: str) -> List[List]:
        """
        Read content from a CSV file.
        
        Args:
            filepath (str): The path to the CSV file to read from
            
        Returns:
            List[List]: The content of the CSV file as a list of lists, where each inner list represents a row
            
        Raises:
            FileNotFoundError: If the file does not exist
            Exception: If there is an error reading the file
        """
        try:
            if not os.path.exists(filepath):
                raise FileNotFoundError(f"File {filepath} does not exist")
            
            df = pd.read_csv(filepath)
            # Convert DataFrame to list of lists
            return [df.columns.tolist()] + df.values.tolist()
        except Exception as e:
            print(f"Error reading CSV file: {e}")
            raise

    @staticmethod
    def read_excel(filepath: str) -> Dict:
        """
        Read content from an Excel file.
        
        Args:
            filepath (str): The path to the Excel file to read from
            
        Returns:
            Dict: A dictionary with sheet names as keys and their data as values (list of lists)
            
        Raises:
            FileNotFoundError: If the file does not exist
            Exception: If there is an error reading the file
        """
        try:
            if not os.path.exists(filepath):
                raise FileNotFoundError(f"File {filepath} does not exist")
            
            excel_file = pd.ExcelFile(filepath)
            result = {}
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(filepath, sheet_name=sheet_name)
                # Convert DataFrame to list of lists
                result[sheet_name] = [df.columns.tolist()] + df.values.tolist()
                
            return result
        except Exception as e:
            print(f"Error reading Excel file: {e}")
            raise
    
    def read(self, filepath: str, file_type: Optional[str] = None) -> Union[str, List[List], Dict]:
        """
        Read content from a file, automatically determining the format based on file extension.
        
        Args:
            filepath (str): The path to the file to read from
            file_type (str, optional): The type of file to read ('markdown', 'word', 'csv', 'excel' or None to auto-detect)
            
        Returns:
            Union[str, List[List], Dict]: The content of the file in a format depending on the file type
            
        Raises:
            ValueError: If unable to determine file type or unsupported file type
            FileNotFoundError: If the file does not exist
            Exception: If there is an error reading the file
        """
        # Determine file type if not explicitly provided
        if file_type is None:
            if filepath.endswith('.md') or filepath.endswith('.markdown'):
                file_type = 'markdown'
            elif filepath.endswith('.docx'):
                file_type = 'word'
            elif filepath.endswith('.csv'):
                file_type = 'csv'
            elif filepath.endswith('.xlsx') or filepath.endswith('.xls'):
                file_type = 'excel'
            else:
                raise ValueError("Unable to determine file type from extension. Please specify file_type.")
        
        if file_type.lower() in ['md', 'markdown']:
            return self.read_markdown(filepath)
        elif file_type.lower() in ['docx', 'word']:
            return self.read_word(filepath)
        elif file_type.lower() in ['csv']:
            return self.read_csv(filepath)
        elif file_type.lower() in ['excel', 'xlsx', 'xls']:
            return self.read_excel(filepath)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")


    @staticmethod
    def write_markdown(content: str, filepath: str, overwrite: bool = False) -> bool:
        """
        Write content to a Markdown file.
        
        Args:
            content (str): The content to write to the file
            filepath (str): The path to the file to write to
            overwrite (bool): Whether to overwrite the file if it exists
            
        Returns:
            bool: True if successful, False otherwise
        """
        try:
            # Ensure the directory exists
            os.makedirs(os.path.dirname(filepath) if os.path.dirname(filepath) else '.', exist_ok=True)
            
            # Check if file exists and handle overwrite
            if os.path.exists(filepath) and not overwrite:
                raise FileExistsError(f"File {filepath} already exists and overwrite is set to False")
            
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(content)
            return True
        except Exception as e:
            print(f"Error writing to markdown file: {e}")
            return False

    @staticmethod
    def write_word(content: str, filepath: str, overwrite: bool = False) -> bool:
        """
        Write content to a Word file.
        
        Args:
            content (str): The content to write to the file
            filepath (str): The path to the file to write to
            overwrite (bool): Whether to overwrite the file if it exists
            
        Returns:
            bool: True if successful, False otherwise
        """
        try:
            # Ensure the directory exists
            os.makedirs(os.path.dirname(filepath) if os.path.dirname(filepath) else '.', exist_ok=True)
            
            # Check if file exists and handle overwrite
            if os.path.exists(filepath) and not overwrite:
                raise FileExistsError(f"File {filepath} already exists and overwrite is set to False")
            
            doc = Document()
            
            # Split content into paragraphs and add them
            paragraphs = content.split('\n\n')
            for para in paragraphs:
                if para.startswith('#'):
                    # Simple markdown heading support
                    level = len(para) - len(para.lstrip('#'))
                    doc.add_heading(para.lstrip('# ').lstrip(), level=min(level, 6))
                else:
                    doc.add_paragraph(para)
            
            doc.save(filepath)
            return True
        except Exception as e:
            print(f"Error writing to Word file: {e}")
            return False

    @staticmethod
    def write_csv(data: Union[List[List], List[Dict]], filepath: str, overwrite: bool = False, **kwargs) -> bool:
        """
        Write data to a CSV file.
        
        Args:
            data (Union[List[List], List[Dict]]): The data to write to the CSV file. 
                                                 Can be a list of lists (rows) or list of dictionaries.
            filepath (str): The path to the file to write to
            overwrite (bool): Whether to overwrite the file if it exists
            **kwargs: Additional arguments to pass to pandas.DataFrame.to_csv()
                      For example: index, header, sep, etc.
            
        Returns:
            bool: True if successful, False otherwise
            
        Examples:
            >>> data = [['Name', 'Age'], ['Alice', 30], ['Bob', 25]]
            >>> DocumentWriter.write_csv(data, 'people.csv')
            True
            
            >>> data = [{'Name': 'Alice', 'Age': 30}, {'Name': 'Bob', 'Age': 25}]
            >>> DocumentWriter.write_csv(data, 'people.csv', index=False)
            True
        """
        try:
            # Ensure the directory exists
            os.makedirs(os.path.dirname(filepath) if os.path.dirname(filepath) else '.', exist_ok=True)
            
            # Check if file exists and handle overwrite
            if os.path.exists(filepath) and not overwrite:
                raise FileExistsError(f"File {filepath} already exists and overwrite is set to False")
            
            # Convert data to DataFrame
            df = pd.DataFrame(data)
            
            # Write to CSV with default parameters
            csv_kwargs = {}
            csv_kwargs.update(kwargs)
            df.to_csv(filepath, **csv_kwargs)
            return True
        except Exception as e:
            print(f"Error writing to CSV file: {e}")
            return False

    @staticmethod
    def write_excel(data: Union[List[List], List[Dict]], filepath: str, sheet_name: str = 'Sheet1', 
                    overwrite: bool = False, **kwargs) -> bool:
        """
        Write data to an Excel file.
        
        Args:
            data (Union[List[List], List[Dict]]): The data to write to the Excel file. 
                                                 Can be a list of lists (rows) or list of dictionaries.
            filepath (str): The path to the file to write to
            sheet_name (str): The name of the sheet to write to. Defaults to 'Sheet1'.
            overwrite (bool): Whether to overwrite the file if it exists
            **kwargs: Additional arguments to pass to pandas.DataFrame.to_excel()
                      For example: index, header, etc.
            
        Returns:
            bool: True if successful, False otherwise
            
        Examples:
            >>> data = [['Name', 'Age'], ['Alice', 30], ['Bob', 25]]
            >>> DocumentWriter.write_excel(data, 'people.xlsx')
            True
            
            >>> data = [{'Name': 'Alice', 'Age': 30}, {'Name': 'Bob', 'Age': 25}]
            >>> DocumentWriter.write_excel(data, 'people.xlsx', sheet_name='Employees', index=False)
            True
        """
        try:
            # Ensure the directory exists
            os.makedirs(os.path.dirname(filepath) if os.path.dirname(filepath) else '.', exist_ok=True)
            
            # Check if file exists and handle overwrite
            if os.path.exists(filepath) and not overwrite:
                raise FileExistsError(f"File {filepath} already exists and overwrite is set to False")
            
            # Convert data to DataFrame
            df = pd.DataFrame(data)
            
            # Write to Excel with default parameters
            excel_kwargs = {}
            excel_kwargs.update(kwargs)
            
            with pd.ExcelWriter(filepath, engine='openpyxl') as writer:
                df.to_excel(writer, sheet_name=sheet_name, **excel_kwargs)
            return True
        except Exception as e:
            print(f"Error writing to Excel file: {e}")
            return False

    def write(self, content: Union[str, List[List], List[Dict]], filepath: str, file_type: Optional[str] = None, overwrite: bool = False,
              **kwargs) -> bool:
        """
        Write content to a file, automatically determining the format based on file extension.
        
        Args:
            content (Union[str, List[List], List[Dict]]): The content to write to the file. 
                                                         For markdown and word files, this should be a string.
                                                         For csv and excel files, this can be a list of lists or list of dicts.
            filepath (str): The path to the file to write to
            file_type (str, optional): The type of file to write ('markdown', 'word', 'csv', 'excel' or None to auto-detect)
            overwrite (bool): Whether to overwrite the file if it exists
            **kwargs: Additional arguments for specific file types
            
        Returns:
            bool: True if successful, False otherwise
        """
        # Determine file type if not explicitly provided
        if file_type is None:
            if filepath.endswith('.md') or filepath.endswith('.markdown'):
                file_type = 'markdown'
            elif filepath.endswith('.docx'):
                file_type = 'word'
            elif filepath.endswith('.csv'):
                file_type = 'csv'
            elif filepath.endswith('.xlsx') or filepath.endswith('.xls'):
                file_type = 'excel'
            else:
                raise ValueError("Unable to determine file type from extension. Please specify file_type.")
        
        if file_type.lower() in ['md', 'markdown']:
            # Ensure content is string for markdown
            if not isinstance(content, str):
                raise TypeError("Content must be a string for markdown files")
            return self.write_markdown(content, filepath, overwrite)
        elif file_type.lower() in ['docx', 'word']:
            # Ensure content is string for word
            if not isinstance(content, str):
                raise TypeError("Content must be a string for Word files")
            return self.write_word(content, filepath, overwrite)
        elif file_type.lower() in ['csv']:
            # For CSV, we expect data to be passed in kwargs or content
            data = kwargs.get('data', content)
            return self.write_csv(data, filepath, overwrite, **kwargs)
        elif file_type.lower() in ['excel', 'xlsx', 'xls']:
            # For Excel, we expect data to be passed in kwargs or content
            data = kwargs.get('data', content)
            sheet_name = kwargs.get('sheet_name', 'Sheet1')
            return self.write_excel(data, filepath, sheet_name, overwrite, **kwargs)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")


doc_handler = DocumentHandler()

